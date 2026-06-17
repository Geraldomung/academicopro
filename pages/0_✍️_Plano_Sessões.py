import streamlit as st
import os
from dotenv import load_dotenv
import docx
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import google.generativeai as genai
from datetime import date, datetime
import pandas as pd
import sqlite3
import json
import re

# --- Configurações Iniciais ---
load_dotenv()

# --- Configuração do Gemini Client ---
def setup_gemini_client():
    """Configura e retorna o modelo Gemini."""
    google_api_key = os.getenv("GOOGLE_API_KEY") or st.secrets.get("GOOGLE_API_KEY")
    if not google_api_key:
        st.error("ERRO: Chave da API do Google Gemini não encontrada.")
        st.stop()
    
    genai.configure(api_key=google_api_key)
    try:
        # Usamos o gemini-2.5-flash para respostas rápidas e estruturadas (JSON)
        return genai.GenerativeModel('gemini-2.5-flash')
    except Exception as e:
        st.error(f"Erro ao configurar o modelo Gemini: {e}")
        st.stop()

gemini_model = setup_gemini_client()

# --- Configuração da Página Streamlit ---
st.set_page_config(
    page_title="Gerador de Plano de Sessão (Moçambique)",
    layout="wide",
    page_icon="🇲🇿"
)

# --- Funções de Banco de Dados (SQLite) ---
def init_db():
    conn = sqlite3.connect('db/planos_sessao.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS registros_sessao
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  data_hora TEXT NOT NULL,
                  ucp TEXT,
                  unidade_aprendizagem TEXT,
                  formador TEXT,
                  instituto TEXT,
                  duracao_minutos INTEGER)''')
    conn.commit()
    conn.close()

def registrar_sessao(ucp, ua, formador, instituto, duracao_minutos):
    conn = None
    try:
        conn = sqlite3.connect('db/planos_sessao.db')
        c = conn.cursor()
        data_hora = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        c.execute('''INSERT INTO registros_sessao
                     (data_hora, ucp, unidade_aprendizagem, formador, instituto, duracao_minutos)
                     VALUES (?, ?, ?, ?, ?, ?)''',
                  (data_hora, ucp, ua, formador, instituto, duracao_minutos))
        conn.commit()
    except sqlite3.Error as e:
        st.error(f"Erro ao registrar o plano de sessão: {e}")
    finally:
        if conn:
            conn.close()

def get_historico_sessoes():
    conn = sqlite3.connect('db/planos_sessao.db')
    df = pd.read_sql_query("SELECT * FROM registros_sessao ORDER BY data_hora DESC", conn)
    conn.close()
    return df

init_db()

# --- Componentes de Competências ---
COMPETENCIAS = {
    "C": "Cognitivo (Conhecimentos)",
    "H": "Habilidades (Psicomotor)",
    "A": "Atitudes (Afetivo)"
}

# --- Interface do Usuário Streamlit ---
st.title("🇲🇿 Gerador de Planos de Sessão (Ensino Técnico)")
st.write("Crie planos de sessão detalhados, baseados no modelo de Moçambique, utilizando IA.")

tab1, tab2 = st.tabs(["Gerar Novo Plano", "Histórico de Planos"])

with tab1:
    # --- Dados de Identificação ---
    st.subheader("📝 Informações de Identificação")
    col1, col2 = st.columns(2)
    with col1:
        instituto = st.text_input("Nome do Instituto:", placeholder="Ex: IIC Eduardo Mondlane", key="instituto")
        ucp = st.text_input("Unidade Curricular (UCP):", placeholder="Ex: Contabilidade Geral", key="ucp")
        codigo = st.text_input("Código da UCP/UCAD:", placeholder="Ex: UCADG097002", key="codigo")
    with col2:
        formador = st.text_input("Nome do Formador(a):", placeholder="Ex: Hagira E. M. Nhampossa", key="formador")
        data_sessao = st.date_input("Data da Sessão:", value=date.today(), key="data_sessao")
        duracao_minutos = st.number_input("Duração da Sessão (minutos):", min_value=30, value=120, step=30, key="duracao_minutos")

    st.markdown("---")
    
    # --- Componentes de Competência ---
    st.subheader("🎯 Componentes de Competência")
    ua = st.text_area(
        "Unidade da Aprendizagem (Elementos de competência):",
        placeholder="Descreva o elemento de competência ou o tópico central da sessão...",
        key="ua",
        height=100
    )
    
    criterios_desempenho = st.text_area(
        "Competências a desenvolver (Critérios de Desempenho):",
        placeholder="Descreva o que o formando deve ser capaz de fazer ao fim da sessão, um critério por linha...",
        key="criterios_desempenho",
        height=100
    )

    # NOVO CAMPO DE ENTRADA PARA O RA
    resultados_aprendizagem = st.text_area(
        "Resultados de Aprendizagem (RAs) - Predefinidos:",
        placeholder="Digite um ou mais Resultados de Aprendizagem (um por linha)...",
        key="ra_predefinidos",
        height=100
    )
    
    conteudo_detalhado = st.text_area(
        "Conteúdo a ser Abordado (Tópicos Principais):",
        placeholder="Descreva o conteúdo específico da sessão (Ex: O Património: Noção e Elementos)",
        key="conteudo_detalhado",
        height=100
    )
    
    st.markdown("---")

    # --- Lógica de Geração do Plano de Sessão ---
    if st.button("🚀 Gerar Plano de Sessão com IA", type="primary", key="gerar"):
        # ATUALIZAÇÃO: Adicionar resultados_aprendizagem à verificação
        if not all([instituto, ucp, ua, criterios_desempenho, conteudo_detalhado, resultados_aprendizagem]):
            st.warning("Por favor, preencha todas as informações essenciais (Instituto, UCP, Unidade de Aprendizagem, Critérios, Conteúdo e **Resultados de Aprendizagem**).")
            st.stop()
        
        with st.spinner("Gerando plano de sessão estruturado..."):
            try:
                # O PROMPT - Essencial para obter o JSON estruturado
                prompt = f"""
                Você é um especialista em desenvolvimento curricular para o Ensino Técnico Profissional (ETP) em Moçambique. Crie um Plano de Sessão detalhado no formato JSON.

                O objeto JSON raiz deve conter as seguintes chaves de nível superior: "ra_usado", "tabela_didactica", "observacoes", e "recursos_gerais".

                Dados de entrada:
                Instituto: {instituto}
                UCP: {ucp}
                Unidade da Aprendizagem (UA): {ua}
                Critérios de Desempenho: {criterios_desempenho}
                Conteúdo da Sessão: {conteudo_detalhado}
                Resultados de Aprendizagem (RAs) predefinidos: {resultados_aprendizagem}
                Duração da Sessão: {duracao_minutos} minutos.

                1.  **"ra_usado"**: Repita os Resultados de Aprendizagem (RAs) fornecidos pelo professor. Este campo serve apenas como registro.

                2.  **"recursos_gerais"**: Liste 3-5 recursos didáticos gerais (Ex: Manuais de Contabilidade, Projetor, Quadro Branco, etc.).

                3.  **"observacoes"**: Gere 3-5 pontos de observação pedagógica e logística (Ex: Métodos a serem usados, Inclusão, Ética, Prevenção de COVID-19).

                4.  **"tabela_didactica"**: Crie uma lista de objetos JSON que detalhe os momentos da sessão (Início/Revisão, Desenvolvimento, Consolidação/Avaliação). Você **DEVE** usar os Resultados de Aprendizagem (RAs) **FORNECIDOS PELO PROFESSOR** como o elemento central de cada objeto. Para cada RA fornecido, crie pelo menos um (ou mais) objetos na lista para detalhar a sua execução didática. Cada objeto deve ter as seguintes chaves obrigatórias e na ordem correta, espelhando a estrutura do Plano de Sessão:

                    * **"ra"**: **USE O TEXTO COMPLETO DE UM DOS RAs FORNECIDOS (ou uma parte, se for muito longo, desde que mantenha a ideia principal)**. Se o professor forneceu múltiplos RAs, crie uma sequência de objetos para cobrir todos eles.
                    * **"competencias_c"**: Booleano (true/false) para a componente Cognitiva (C)
                    * **"competencias_h"**: Booleano (true/false) para a componente Habilidades (H)
                    * **"competencias_a"**: Booleano (true/false) para a componente Atitudes (A)
                    * **"conteudo_especifico"**: O tópico específico dentro do Conteúdo Geral que será coberto para atingir este RA.
                    * **"metodo_tecnica"**: O Método/Técnica de ensino para este momento.
                    * **"atividade_formador"**: O que o Formador faz neste momento.
                    * **"atividade_formando"**: O que o Formando faz neste momento.
                    * **"tempo_estimado_min"**: O tempo em minutos para este momento (A soma total deve ser coerente com a Duração da Sessão fornecida).
                    * **"criterios_avaliacao"**: Quais critérios serão usados para avaliar este RA.
                    * **"evidencia"**: A evidência de aprendizagem (Ex: Oral, Escrita, Prática, Observação).
                    * **"tecnicas_instrumento"**: A Técnica/Instrumento de avaliação (Ex: Avaliação Formativa/ Ficha de verificação).
                """
                
                response = gemini_model.generate_content(prompt)
                plano_gerado = response.text
                
                # Robust JSON parsing
                json_match = re.search(r"```json\s*(\{.*\})\s*```", plano_gerado, re.DOTALL)
                if not json_match:
                    json_match = re.search(r"(\{.*\})", plano_gerado, re.DOTALL)
                
                if json_match:
                    json_string = json_match.group(1)
                    plano_json = json.loads(json_string)
                    
                    # Convert to DataFrame for display and easy handling
                    tabela_data = plano_json.get("tabela_didactica", [])
                    df_tabela = pd.DataFrame(tabela_data)
                    
                    # Adicionar colunas faltantes para garantir que o DataFrame seja robusto
                    expected_keys = [
                        "ra", "competencias_c", "competencias_h", "competencias_a", 
                        "conteudo_especifico", "metodo_tecnica", "atividade_formador", 
                        "atividade_formando", "tempo_estimado_min", "criterios_avaliacao", 
                        "evidencia", "tecnicas_instrumento"
                    ]

                    for key in expected_keys:
                        if key not in df_tabela.columns:
                            # Preenche colunas faltantes com string vazia
                            df_tabela[key] = '' 
                            
                    # Validate time structure
                    if 'tempo_estimado_min' in df_tabela.columns:
                        try:
                            df_tabela['tempo_estimado_min'] = pd.to_numeric(df_tabela['tempo_estimado_min'], errors='coerce').fillna(0).astype(int)
                        except:
                            st.warning("Erro ao converter tempos para numérico. O cálculo total pode estar incorreto.")
                            df_tabela['tempo_estimado_min'] = 0

                    
                    # Store data in session state
                    st.session_state.plano_sessao_json = plano_json
                    st.session_state.df_tabela = df_tabela
                    
                    # Registrar no banco de dados
                    registrar_sessao(ucp, ua, formador, instituto, duracao_minutos)
                    
                    st.success("Plano de Sessão gerado com sucesso!")
                        
                else:
                    st.error("Não foi possível extrair um JSON válido da resposta da IA. Tente novamente ou ajuste os dados de entrada.")
                    st.code(plano_gerado)
                    st.stop()
                    
            except Exception as e:
                st.error(f"Erro inesperado ao gerar o plano de sessão: {str(e)}")
                st.stop()

    # --- Exibição dos Resultados ---
    if 'plano_sessao_json' in st.session_state and 'df_tabela' in st.session_state:
        st.subheader("📋 Plano de Sessão Gerado")

        # Exibição dos Dados Gerais
        col_display1, col_display2, col_display3 = st.columns(3)
        with col_display1:
            st.markdown(f"**UCP:** {ucp}")
            st.markdown(f"**Instituto:** {instituto}")
        with col_display2:
            st.markdown(f"**Formador:** {formador}")
            st.markdown(f"**Data da Sessão:** {data_sessao.strftime('%d/%m/%Y')}")
        with col_display3:
            st.markdown(f"**UA:** {ua}")
            st.markdown(f"**Duração:** {duracao_minutos} min")
        
        st.markdown("---")

        # NOVO - Exibição dos RAs fornecidos
        st.subheader("💡 Resultados de Aprendizagem (Predefinidos)")
        st.text_area(label="Resultados de Aprendizagem (RAs) Predefinidos:", value=resultados_aprendizagem, height=100, disabled=True)
        
        st.markdown("---")
        
        # Tabela Didática
        st.subheader("Estrutura Didática por Resultado de Aprendizagem")
        
        # Prepara a tabela para exibição no Streamlit
        df_display = st.session_state.df_tabela.copy()
        
        # Adiciona colunas de Competências (C/H/A) como checkmarks
        C_COL = COMPETENCIAS['C']
        H_COL = COMPETENCIAS['H']
        A_COL = COMPETENCIAS['A']

        df_display[C_COL] = df_display['competencias_c'].apply(lambda x: 'X' if x in [True, 'True', 1, 'X'] else '')
        df_display[H_COL] = df_display['competencias_h'].apply(lambda x: 'X' if x in [True, 'True', 1, 'X'] else '')
        df_display[A_COL] = df_display['competencias_a'].apply(lambda x: 'X' if x in [True, 'True', 1, 'X'] else '')
        
        # Renomear as colunas
        df_display = df_display.rename(columns={
            'ra': 'RA',
            'conteudo_especifico': 'Conteúdo Específico',
            'metodo_tecnica': 'Método/Técnica',
            'atividade_formador': 'Atividade Formador',
            'atividade_formando': 'Atividade Formando',
            'tempo_estimado_min': 'Tempo (min)',
            'criterios_avaliacao': 'Critérios de Avaliação',
            'evidencia': 'Evidência',
            'tecnicas_instrumento': 'Técnica/Instrumento'
        })
        
        # Definir as colunas para exibição usando os nomes RENOMEADOS
        final_display_cols = [
            'RA', C_COL, H_COL, A_COL,
            'Conteúdo Específico', 'Método/Técnica', 'Atividade Formador', 
            'Atividade Formando', 'Tempo (min)', 
            'Critérios de Avaliação', 'Evidência', 'Técnica/Instrumento'
        ]
        
        # Selecionar as colunas usando os nomes RENOMEADOS
        df_display = df_display[final_display_cols]

        st.dataframe(df_display, use_container_width=True)
        st.markdown(f"**Tempo Total Estimado:** {st.session_state.df_tabela['tempo_estimado_min'].sum()} minutos (Target: {duracao_minutos} min)")

        st.markdown("---")

        # Observações e Recursos
        col_obs, col_rec = st.columns(2)
        with col_obs:
            st.subheader("⚠️ Observações")
            observacoes = st.session_state.plano_sessao_json.get('observacoes', [])
            if observacoes:
                for obs in observacoes:
                    st.markdown(f"- {obs}")
        
        with col_rec:
            st.subheader("📚 Recursos Gerais")
            recursos = st.session_state.plano_sessao_json.get('recursos_gerais', [])
            if recursos:
                for rec in recursos:
                    st.markdown(f"- {rec}")
        
        # --- Função para Criar Documento Word ---
        def criar_documento_word_sessao():
            doc = docx.Document()
            
            # Título do Documento
            doc.add_heading('PLANO DE SESSÃO', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Gerado por IA para ETP (Moçambique) em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            
            # Tabela de Identificação (Seguindo o Padrão do Exemplo)
            table_id = doc.add_table(rows=5, cols=4)
            table_id.style = 'Table Grid'
            
            # Dados de Identificação
            table_id.cell(0, 0).merge(table_id.cell(0, 3)).text = f"Instituto: {instituto}  |  UCP: {ucp}"
            table_id.cell(1, 0).text = f"CÓDIGO: {codigo}"
            table_id.cell(1, 1).text = f"DURAÇÃO DA SESSÃO: {duracao_minutos} min"
            table_id.cell(1, 2).text = f"DATA: {data_sessao.strftime('%d-%m-%Y')}"
            table_id.cell(1, 3).text = f"HORÁRIO: {datetime.now().strftime('%H:%M')}" # Hora da Geração

            table_id.cell(2, 0).merge(table_id.cell(2, 3)).text = f"Unidade da Aprendizagem (Elementos de competência): {ua}"
            table_id.cell(3, 0).merge(table_id.cell(3, 3)).text = f'Competências a desenvolver (Critérios de desempenho): "{criterios_desempenho}"'
            
            # ADICIONAR RAs FORNECIDOS
            table_id.cell(4, 0).merge(table_id.cell(4, 3)).text = f'Resultados de Aprendizagem: {resultados_aprendizagem}'
            
            # Cabeçalhos da Tabela Didática
            doc.add_page_break()
            doc.add_heading('Estrutura Didática Detalhada', level=1)
            
            # Tabela em Landscape (para o plano detalhado)
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11.69)
            section.page_height = Inches(8.27)
            
            df_word = st.session_state.df_tabela.copy()
            
            # Colunas da tabela final no Word
            final_columns = [
                'ra', 'competencias_c', 'competencias_h', 'competencias_a', 
                'conteudo_especifico', 'metodo_tecnica', 'atividade_formador', 
                'atividade_formando', 'tempo_estimado_min', 'criterios_avaliacao', 
                'evidencia', 'tecnicas_instrumento'
            ]
            
            table = doc.add_table(rows=1, cols=len(final_columns))
            table.style = 'Table Grid'
            
            # Cabeçalhos da Tabela
            hdr_cells = table.rows[0].cells
            headers_map = {
                'ra': 'Resultado da Aprendizagem',
                'competencias_c': 'C', 'competencias_h': 'H', 'competencias_a': 'A',
                'conteudo_especifico': 'Conteúdo',
                'metodo_tecnica': 'Método/Técnica',
                'atividade_formador': 'Atividade (Formador)',
                'atividade_formando': 'Atividade (Formando)',
                'tempo_estimado_min': 'Tempo (min)',
                'criterios_avaliacao': 'Critérios de Avaliação',
                'evidencia': 'Evidência',
                'tecnicas_instrumento': 'Técnicas/Instrumento'
            }

            for i, col_key in enumerate(final_columns):
                hdr_cells[i].text = headers_map.get(col_key, col_key)
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Preenche tabela com dados do DataFrame
            for idx, row in df_word.iterrows():
                row_cells = table.add_row().cells
                for i, col_key in enumerate(final_columns):
                    value = row.get(col_key, '')
                    
                    if col_key in ['competencias_c', 'competencias_h', 'competencias_a']:
                        # Usa 'X' se for True, 'True', 1 ou 'X'
                        row_cells[i].text = 'X' if value in [True, 'True', 1, 'X'] else '' 
                    elif col_key == 'tempo_estimado_min':
                        row_cells[i].text = str(value)
                    else:
                        row_cells[i].text = str(value)

            # Voltar para Retrato (Observações e Assinaturas)
            doc.add_page_break()
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            section.orientation = WD_ORIENT.PORTRAIT

            # Observações
            doc.add_heading('Observações', level=2)
            observacoes_doc = st.session_state.plano_sessao_json.get('observacoes', [])
            for obs in observacoes_doc:
                doc.add_paragraph(f"• {obs}")
            doc.add_paragraph()

            # Recursos
            doc.add_heading('Recursos', level=2)
            recursos_doc = st.session_state.plano_sessao_json.get('recursos_gerais', [])
            for rec in recursos_doc:
                doc.add_paragraph(f"• {rec}")
            doc.add_paragraph()
            
            # Seção de Assinaturas (Espaçamento)
            doc.add_heading('Aprovação e Assinaturas', level=2)
            doc.add_paragraph(f"Formador(a): {formador}")
            doc.add_paragraph(f"Data Apresentação: {date.today().strftime('%d/%m/%Y')}")
            doc.add_paragraph("Assinatura: _________________________")
            doc.add_paragraph()
            doc.add_paragraph("Supervisor:")
            doc.add_paragraph("Assinatura: _________________________")

            # Salva em bytes
            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)
            return bio
        
        # Botão de download
        st.download_button(
            label="📥 Baixar Plano de Sessão (Word)",
            data=criar_documento_word_sessao(),
            file_name=f"Plano_Sessao_{ucp.replace(' ', '_')}_{data_sessao.strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.markdown("---")


with tab2:
    st.subheader("📊 Histórico de Planos de Sessão Gerados")
    historico_df = get_historico_sessoes()
    if not historico_df.empty:
        historico_df['data_hora'] = pd.to_datetime(historico_df['data_hora']).dt.strftime('%d/%m/%Y %H:%M:%S')
        st.dataframe(historico_df, use_container_width=True)
    else:
        st.info("Nenhum plano de sessão foi registrado ainda.")

# Rodapé
st.markdown("""
    <div style="text-align: center; margin-top: 30px; color: #666;">
        <p style="font-size: 0.8em;">Gerador de Plano de Sessão para ETP - Moçambique</p>
    </div>
    """, unsafe_allow_html=True)