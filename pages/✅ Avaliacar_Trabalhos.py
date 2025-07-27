import streamlit as st
from docx import Document
import PyPDF2
import sqlite3
import re
import os
from io import BytesIO
import google.generativeai as genai
from difflib import SequenceMatcher
import spacy
from datetime import datetime
from collections import Counter
import requests
from typing import Dict, List, Tuple, Optional
from requests.exceptions import RequestException
from dotenv import load_dotenv

# Configurações iniciais
st.set_page_config(layout="wide", page_title="Avaliador Acadêmico Completo")
nlp = spacy.load("pt_core_news_sm")  # Carrega modelo em português

# --- Constantes ---
PLAGIARISM_API_KEY = "4ANN8MjBicA3xEPcmqQ5AHyHKsH9DsfS"
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY") or st.secrets.get("GOOGLE_API_KEY")

if not GOOGLE_API_KEY:
    st.error("ERRO: Chave GOOGLE_API_KEY não encontrada.")
    st.stop()

try:
    genai.configure(api_key=GOOGLE_API_KEY)
    model = genai.GenerativeModel('gemini-2.5-flash')
except Exception as e:
    st.error(f"Erro ao inicializar o modelo Gemini: {str(e)}")
    st.stop()

PALAVRAS_VAZIAS = {
    'o', 'a', 'os', 'as', 'um', 'uns', 'uma', 'umas', 'de', 'do', 'da', 'dos', 'das',
    'em', 'no', 'na', 'nos', 'nas', 'por', 'pelo', 'pela', 'pelos', 'pelas',
    'para', 'ao', 'à', 'aos', 'às', 'com', 'sem', 'sob', 'sobre', 'entre',
    'que', 'e', 'ou', 'se', 'mas', 'como', 'quando', 'onde', 'porque'
}

# Padrões de citação para ABNT e APA
NORMAS_ACADEMICAS = {
    "ABNT": {
        "citacao_direta": r'^".+"\s\([A-ZÀ-Ü][a-zà-ü]+, \d{4}, p\. \d+\)$',
        "citacao_indireta": r'^\([A-ZÀ-Ü][a-zà-ü]+(?:; [A-ZÀ-Ü][a-zà-ü]+)*, \d{4}\)$',
        "referencia_livro": r'^[A-ZÀ-Ü][a-zà-ü]+, [A-ZÀ-Ü]\. \w+\. .+?: .+?, \d{4}\.$',
        "margens": {"esquerda": 3, "direita": 2, "superior": 3, "inferior": 2},  # em cm
    },
    "APA": {
        "citacao_direta": r'^".+"\s\([A-Z][a-z]+, \d{4}, p\. \d+\)$',
        "citacao_indireta": r'^\([A-Z][a-z]+(?: & [A-Z][a-z]+)*, \d{4}\)$',
        "referencia_livro": r'^[A-Z][a-z]+, [A-Z]\. \d{4}\. .+?\. .+?: .+?\.$',
        "margens": {"esquerda": 2.54, "direita": 2.54, "superior": 2.54, "inferior": 2.54},  # em cm
    },
}

ESTRUTURAS_DOCUMENTOS = {
    "TCC": {
        "ABNT": {
            "secoes_esperadas": ["Introdução", "Revisão da Literatura", "Metodologia", "Resultados", "Conclusão"],
            "extensao_media": 10000,
            "citacoes_minimas": 20
        },
        "APA": {
            "secoes_esperadas": ["Introduction", "Literature Review", "Methods", "Results", "Discussion"],
            "extensao_media": 10000,
            "citacoes_minimas": 25
        }
    },
    "Artigo": {
        "ABNT": {
            "secoes_esperadas": ["Resumo", "Introdução", "Métodos", "Resultados", "Discussão"],
            "extensao_media": 5000,
            "citacoes_minimas": 15
        },
        "APA": {
            "secoes_esperadas": ["Abstract", "Introduction", "Methods", "Results", "Discussion"],
            "extensao_media": 5000,
            "citacoes_minimas": 18
        }
    },
    "Dissertação": {
        "ABNT": {
            "secoes_esperadas": ["Capítulo 1", "Capítulo 2", "Capítulo 3", "Capítulo 4", "Conclusões"],
            "extensao_media": 30000,
            "citacoes_minimas": 50
        },
        "APA": {
            "secoes_esperadas": ["Chapter 1", "Chapter 2", "Chapter 3", "Chapter 4", "Conclusions"],
            "extensao_media": 30000,
            "citacoes_minimas": 60
        }
    },
    "Outro": {
        "ABNT": {
            "secoes_esperadas": [],
            "extensao_media": 1000,
            "citacoes_minimas": 10
        },
        "APA": {
            "secoes_esperadas": [],
            "extensao_media": 1000,
            "citacoes_minimas": 12
        }
    }
}

# --- Banco de Dados ---
def init_db():
    conn = sqlite3.connect('projetos_academicos.db')
    c = conn.cursor()
    
    c.execute('''CREATE TABLE IF NOT EXISTS projetos
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                 nome TEXT UNIQUE,
                 secoes TEXT,
                 avaliacao TEXT,
                 tipo_documento TEXT,
                 norma_academica TEXT,
                 data TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS historico
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                 projeto_id INTEGER,
                 acao TEXT,
                 data TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                 FOREIGN KEY(projeto_id) REFERENCES projetos(id))''')
    
    conn.commit()
    conn.close()

def save_to_db(nome_projeto: str, secoes: Dict, avaliacao: str, tipo_doc: str, norma: str) -> bool:
    conn = sqlite3.connect('projetos_academicos.db')
    c = conn.cursor()
    try:
        c.execute('''INSERT OR REPLACE INTO projetos (nome, secoes, avaliacao, tipo_documento, norma_academica)
                     VALUES (?, ?, ?, ?, ?)''', 
                  (nome_projeto, str(secoes), avaliacao, tipo_doc, norma))
        conn.commit()
        return True
    except sqlite3.Error as e:
        st.error(f"Erro ao salvar no banco de dados: {e}")
        return False
    finally:
        conn.close()

def load_from_db() -> List[Tuple]:
    conn = sqlite3.connect('projetos_academicos.db')
    c = conn.cursor()
    c.execute("SELECT id, nome, tipo_documento, norma_academica, data FROM projetos ORDER BY data DESC")
    projetos = c.fetchall()
    conn.close()
    return projetos

def get_project_by_id(project_id: int) -> Tuple:
    conn = sqlite3.connect('projetos_academicos.db')
    c = conn.cursor()
    c.execute("SELECT nome, secoes, avaliacao, tipo_documento, norma_academica FROM projetos WHERE id = ?", (project_id,))
    projeto = c.fetchone()
    conn.close()
    return projeto

def delete_project(project_id: int):
    conn = sqlite3.connect('projetos_academicos.db')
    c = conn.cursor()
    c.execute("DELETE FROM projetos WHERE id = ?", (project_id,))
    conn.commit()
    conn.close()

# --- Processamento de Documentos ---
def extract_sections_from_docx(documento) -> Dict[str, str]:
    doc = Document(documento)
    sections = {}
    current_heading = None
    
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            current_heading = paragraph.text
            sections[current_heading] = ""
        elif current_heading:
            sections[current_heading] += paragraph.text + "\n"
    
    return sections

def extract_sections_from_pdf(documento) -> Dict[str, str]:
    reader = PyPDF2.PdfReader(documento)
    text = "\n".join([page.extract_text() for page in reader.pages])
    
    matches = re.finditer(r'(\n\d+[\s\.\-]\s*.+?\n)', text, re.IGNORECASE)
    headings_pos = [(m.group(1).strip(), m.start()) for m in matches]
    
    sections = {}
    for i, (heading, pos) in enumerate(headings_pos):
        end_pos = headings_pos[i+1][1] if i+1 < len(headings_pos) else len(text)
        sections[heading] = text[pos:end_pos].strip()
    
    return sections

def extrair_secoes(documento) -> Dict[str, str]:
    if documento.name.endswith('.docx'):
        return extract_sections_from_docx(documento)
    elif documento.name.endswith('.pdf'):
        return extract_sections_from_pdf(documento)
    else:
        st.error("Formato de arquivo não suportado!")
        return {}

# --- Análise de Texto ---
def contar_palavras_relevantes(texto: str) -> int:
    if not texto.strip():
        return 0
    
    try:
        doc = nlp(texto.lower())
        return sum(
            1 for token in doc 
            if token.is_alpha 
            and token.text not in PALAVRAS_VAZIAS
            and len(token.text) > 2
        )
    except:
        return len([w for w in texto.split() if len(w) > 3])

def analisar_construcao_frasal(texto: str) -> List[str]:
    problemas = []
    doc = nlp(texto)
    frases = list(doc.sents)
    
    if not frases:
        return ["Texto vazio ou não analisável"]
    
    # Comprimento das frases
    comprimentos = [len(frase) for frase in frases]
    if max(comprimentos) > 30:
        problemas.append(f"Frase excessivamente longa ({max(comprimentos)} palavras)")
    
    # Voz passiva
    voz_passiva = sum(1 for token in doc if token.tag_ == "AUX" and token.dep_ == "aux:pass")
    if voz_passiva / len(frases) > 0.3:
        problemas.append(f"Uso excessivo de voz passiva ({voz_passiva} ocorrências)")
    
    # Repetição de palavras
    palavras = [token.text.lower() for token in doc if token.is_alpha and len(token.text) > 3]
    contagem = Counter(palavras)
    repetidas = [p for p, c in contagem.most_common(5) if c > 5]
    if repetidas:
        problemas.append(f"Repetição excessiva: {', '.join(repetidas)}")
    
    return problemas

def detectar_conteudo_ia(texto: str) -> List[str]:
    problemas = []
    padroes_ia = [
        r"\bNo entanto, é importante destacar que\b",
        r"\bVale ressaltar que\b",
        r"\bNesse contexto\b.*\bconstata-se que\b",
        r"\bObserva-se uma\b.*\bde modo a\b",
        r"\bVerifica-se que\b.*\bno sentido de\b"
    ]
    
    for padrao in padroes_ia:
        if re.search(padrao, texto, re.IGNORECASE):
            problemas.append("Padrão típico de texto gerado por IA")
            break
    
    # Análise de diversidade lexical
    palavras = texto.split()
    if len(palavras) > 100:
        repeticoes = sum(1 for i in range(len(palavras)-3) 
                      if palavras[i] == palavras[i+1] == palavras[i+2])
        if repeticoes / len(palavras) > 0.05:
            problemas.append("Baixa diversidade lexical (possível conteúdo de IA)")
    
    return problemas

def verificar_normas(texto: str, tipo_elemento: str) -> List[str]:
    problemas = []
    norma = st.session_state["config"]["norma"]
    padrao = NORMAS_ACADEMICAS[norma].get(tipo_elemento)
    
    if padrao and not re.search(padrao, texto, re.IGNORECASE):
        problemas.append(f"Formatação {tipo_elemento} não conforme {norma}")
    
    return problemas

def verificar_citacoes_completas(texto: str) -> Tuple[int, List[str]]:
    citacoes_encontradas = 0
    problemas = []
    norma = st.session_state["config"]["norma"]
    
    padrao_direta = NORMAS_ACADEMICAS[norma]["citacao_direta"]
    padrao_indireta = NORMAS_ACADEMICAS[norma]["citacao_indireta"]
    
    citacoes_encontradas += len(re.findall(padrao_direta, texto))
    citacoes_encontradas += len(re.findall(padrao_indireta, texto))
    
    # Verifica citações incompletas
    citacoes_incompletas = re.findall(r'\([^)]+\)', texto)
    for cit in citacoes_incompletas:
        if not (re.match(padrao_direta, cit) or re.match(padrao_indireta, cit)):
            problemas.append(f"Citação mal formatada: '{cit}'")

    return citacoes_encontradas, problemas

def verificar_plagio_local(texto: str) -> float:
    """Verificação básica de similaridade local"""
    from difflib import SequenceMatcher
    
    # Base local de trechos conhecidos (pode ser expandida)
    trechos_conhecidos = [
        "A pesquisa qualitativa permite compreender os fenômenos sociais em profundidade.",
        "Segundo autores contemporâneos, a metodologia científica deve ser rigorosa.",
        "Os resultados demonstram uma correlação significativa entre as variáveis analisadas."
    ]
    
    similaridades = [SequenceMatcher(None, texto, trecho).ratio() 
                    for trecho in trechos_conhecidos]
    
    return max(similaridades) if similaridades else 0.0

def verificar_plagio(texto: str) -> Dict:
    """Verifica plágio com tratamento de erros robusto"""
    if not texto.strip() or len(texto.split()) < 20:
        return {"error": "Texto muito curto para análise"}
    
    try:
        # Verifica primeiro se o domínio está resolvível
        try:
            requests.get("https://api.plagiarismcheck.org", timeout=3)
        except RequestException as e:
            return {"error": f"Serviço de plágio indisponível: {str(e)}"}
        
        headers = {
            "Authorization": f"Bearer {PLAGIARISM_API_KEY}",
            "Content-Type": "application/json"
        }
        
        payload = {
            "text": texto[:5000],
            "language": "pt",
            "sensitivity": "medium"
        }
        
        response = requests.post(
            "https://api.plagiarismcheck.org/v1/check",
            headers=headers,
            json=payload,
            timeout=10
        )
        
        if response.status_code == 200:
            return response.json()
        else:
            return {"error": f"Erro na API: {response.status_code}"}
    
    except RequestException as e:
        # Fallback para verificação local básica
        similaridade_local = verificar_plagio_local(texto)
        return {
            "error": f"Falha na conexão com o serviço. Análise local realizada.",
            "score": similaridade_local,
            "local": True
        }
    except Exception as e:
        return {"error": f"Erro inesperado: {str(e)}"}

def analisar_resultado_plagio(resultado: Dict) -> List[str]:
    """Interpreta os resultados da verificação de plágio"""
    problemas = []
    
    if "error" in resultado:
        if resultado.get("local"):
            problemas.append(f"⚠️ {resultado['error']}")
            if resultado['score'] > 0.7:
                problemas.append(f"Alerta: Similaridade local alta ({resultado['score']*100:.1f}%)")
        else:
            problemas.append(f"Erro na verificação: {resultado['error']}")
        return problemas
    
    if resultado.get("score", 0) > 0.25:
        problemas.append(f"⚠️ Possível plágio detectado (similaridade: {resultado['score']*100:.1f}%)")
        
        for match in resultado.get("matches", [])[:3]:
            problemas.append(
                f"- Texto similar em: {match.get('source', 'Fonte desconhecida')} "
                f"(Similaridade: {match.get('score', 0)*100:.1f}%)"
            )
    
    return problemas

def avaliar_secao_completa(texto: str, titulo_secao: str) -> Dict[str, List[str]]:
    resultados = {
        'construcao_frasal': analisar_construcao_frasal(texto),
        'conteudo_ia': detectar_conteudo_ia(texto),
        'citacoes': [],
        'plagio': [],
        'normas': [],
        'geral': []
    }
    
    norma = st.session_state["config"]["norma"]
    tipo_projeto = st.session_state["config"]["tipo_projeto"]
    
    if not texto.strip():
        resultados['geral'].append("Seção vazia")
        return resultados
    
    # Verificação de normas
    if "referência" in titulo_secao.lower() or "references" in titulo_secao.lower():
        resultados['normas'].extend(verificar_normas(texto, "referencia_livro"))
    elif "citação" in titulo_secao.lower() or "citation" in titulo_secao.lower():
        resultados['normas'].extend(verificar_normas(texto, "citacao_direta"))
        resultados['normas'].extend(verificar_normas(texto, "citacao_indireta"))
    
    # Verificação de citações
    citacoes, problemas_cit = verificar_citacoes_completas(texto)
    resultados['citacoes'] = problemas_cit
    
    # Verificação de plágio (apenas para seções com mais de 100 palavras)
    if contar_palavras_relevantes(texto) > 100:
        resultado_plagio = verificar_plagio(texto)
        resultados['plagio'] = analisar_resultado_plagio(resultado_plagio)
    
    # Verificação de densidade de citações
    palavras_relevantes = contar_palavras_relevantes(texto)
    if palavras_relevantes > 0:
        config_doc = ESTRUTURAS_DOCUMENTOS.get(tipo_projeto, ESTRUTURAS_DOCUMENTOS["Outro"])[norma]
        densidade_esperada = config_doc['citacoes_minimas'] / config_doc['extensao_media']
        densidade_atual = citacoes / palavras_relevantes
        
        if densidade_atual < densidade_esperada * 0.5:
            resultados['citacoes'].append(
                f"Baixa densidade de citações (encontradas: {citacoes}, "
                f"esperadas: ~{int(palavras_relevantes * densidade_esperada)})"
            )
    
    # Verificação de estrutura
    estrutura_esperada = config_doc['secoes_esperadas']
    if estrutura_esperada and titulo_secao not in estrutura_esperada:
        resultados['geral'].append(f"Seção '{titulo_secao}' incomum para {tipo_projeto} ({norma})")
    
    return resultados

def gerar_documento_editado(secoes: Dict[str, str]) -> BytesIO:
    doc = Document()
    for titulo, texto in secoes.items():
        doc.add_heading(titulo, level=1)
        doc.add_paragraph(texto)
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def exibir_resultados_avaliacao(resultados: Dict[str, List[str]], titulo: str):
    with st.expander(f"🔍 Análise Detalhada - {titulo}", expanded=False):
        tab1, tab2, tab3, tab4, tab5 = st.tabs(["Construção Frasal", "Originalidade", "Citações", "Plágio", "Normas"])
        
        with tab1:
            if resultados['construcao_frasal']:
                st.warning("Problemas encontrados:")
                for problema in resultados['construcao_frasal']:
                    st.write(f"- {problema}")
            else:
                st.success("Construção frasal adequada")
        
        with tab2:
            if resultados['conteudo_ia']:
                st.error("Possíveis problemas:")
                for problema in resultados['conteudo_ia']:
                    st.write(f"- {problema}")
            else:
                st.success("Estilo original detectado")
        
        with tab3:
            if resultados['citacoes']:
                st.warning("Problemas encontrados:")
                for problema in resultados['citacoes']:
                    st.write(f"- {problema}")
            else:
                st.success("Citações adequadas")
        
        with tab4:
            if resultados['plagio']:
                st.error("Problemas encontrados:")
                for problema in resultados['plagio']:
                    st.write(problema)
            else:
                st.success("Nenhum plágio detectado")
        
        with tab5:
            if resultados['normas']:
                st.error("Problemas de formatação:")
                for problema in resultados['normas']:
                    st.write(f"- {problema}")
            else:
                st.success("Formatação conforme normas")

def exibir_resultados_avaliacao(resultados: Dict[str, List[str]], titulo: str, texto_atual: str):
    with st.expander(f"🔍 Análise Detalhada - {titulo}", expanded=False):
        tab1, tab2, tab3, tab4, tab5 = st.tabs(["Construção Frasal", "Originalidade", "Citações", "Plágio", "Normas"])
        
        with tab1:
            if resultados['construcao_frasal']:
                st.warning("Problemas encontrados:")
                for problema in resultados['construcao_frasal']:
                    st.write(f"- {problema}")
            else:
                st.success("Construção frasal adequada")
        
        with tab2:
            if resultados['conteudo_ia']:
                st.error("Possíveis problemas:")
                for problema in resultados['conteudo_ia']:
                    st.write(f"- {problema}")
            else:
                st.success("Estilo original detectado")
        
        with tab3:
            if resultados['citacoes']:
                st.warning("Problemas encontrados:")
                for problema in resultados['citacoes']:
                    st.write(f"- {problema}")
            else:
                st.success("Citações adequadas")
        
        with tab4:
            if resultados['plagio']:
                st.error("Problemas encontrados:")
                for problema in resultados['plagio']:
                    st.write(problema)
            else:
                st.success("Nenhum plágio detectado")
        
        with tab5:
            if resultados['normas']:
                st.error("Problemas de formatação:")
                for problema in resultados['normas']:
                    st.write(f"- {problema}")
            else:
                st.success("Formatação conforme normas")
        
        # Botão para regenerar texto com base nas observações
        if any(resultados.values()):  # Se houver algum problema detectado
            st.divider()
            if st.button(f"🔄 Regenerar texto da seção '{titulo}'", key=f"regenerate_{titulo}"):
                with st.spinner(f"Reescrevendo a seção '{titulo}'..."):
                    try:
                        # Preparar o prompt com os problemas encontrados
                        problemas_formatados = "\n".join(
                            f"- {problema}" 
                            for categoria in resultados.values() 
                            for problema in categoria
                        )
                        
                        prompt = f"""Reescreva a seguinte seção de um trabalho acadêmico do tipo {st.session_state['tipo_documento']} 
                        seguindo as normas {st.session_state['config']['norma']}. Corrija os problemas identificados:

                        Problemas encontrados:
                        {problemas_formatados}

                        Texto original:
                        {texto_atual}

                        Instruções:
                        - Mantenha o significado original
                        - Corrija todos os problemas apontados
                        - Use linguagem acadêmica apropriada
                        - Não invente novos conteúdos
                        - Retorne apenas o texto reescrito, sem comentários
                        """
                        
                        response = model.generate_content(prompt)
                        st.session_state['secoes'][titulo] = response.text
                        st.rerun()
                        
                    except Exception as e:
                        st.error(f"Erro ao regenerar texto: {e}")

# --- Interface Principal ---
def main():
    st.title("📚 Avaliador Acadêmico Completo")
    init_db()
    
    # --- Configuração Inicial (se ainda não definida) ---
    if "config" not in st.session_state:
        with st.sidebar.expander("⚙️ Configurações do Documento", expanded=True):
            # Selecionar norma (ABNT ou APA)
            norma_selecionada = st.selectbox(
                "📝 Norma Acadêmica",
                ["ABNT", "APA"],
                index=0,
                help="Escolha o estilo de formatação para análise."
            )
            
            # Selecionar tipo de projeto (personalizável)
            tipo_projeto_padrao = st.selectbox(
                "📂 Tipo de Documento",
                ["TCC", "Artigo", "Dissertação", "Outro"],
                index=0,
                help="Selecione o tipo de trabalho para critérios específicos."
            )
            
            # Botão para confirmar
            if st.button("✅ Aplicar Configurações"):
                st.session_state["config"] = {
                    "norma": norma_selecionada,
                    "tipo_projeto": tipo_projeto_padrao,
                }
                st.rerun()
        
        st.stop()  # Pausa até as configurações serem definidas
    
    # Sidebar - Gerenciamento de Projetos
    st.sidebar.header("Projetos Salvos")
    projetos = load_from_db()
    
    if projetos:
        selected_project = st.sidebar.selectbox(
            "Selecione um projeto",
            projetos,
            format_func=lambda x: f"{x[1]} ({x[2]} - {x[3]})"
        )
        
        col1, col2 = st.sidebar.columns(2)
        if col1.button("Carregar Projeto"):
            nome, secoes, avaliacao, tipo_doc, norma = get_project_by_id(selected_project[0])
            st.session_state.update({
                'secoes': eval(secoes),
                'nome_projeto': nome,
                'tipo_documento': tipo_doc,
                'config': {
                    'norma': norma,
                    'tipo_projeto': tipo_doc
                }
            })
            st.rerun()
        
        if col2.button("Excluir Projeto", type="primary"):
            delete_project(selected_project[0])
            st.sidebar.success("Projeto excluído!")
            st.rerun()
    
    # Botão para alterar configurações
    if st.sidebar.button("🔄 Alterar Configurações"):
        del st.session_state["config"]
        st.rerun()
    
    # Área principal
    uploaded_file = st.file_uploader("Carregue seu documento (PDF ou DOCX)", type=["pdf", "docx"])
    
    if uploaded_file and 'secoes' not in st.session_state:
        with st.spinner("Processando documento..."):
            st.session_state['secoes'] = extrair_secoes(uploaded_file)
            st.session_state['nome_projeto'] = os.path.splitext(uploaded_file.name)[0]
            st.session_state['tipo_documento'] = st.session_state["config"]["tipo_projeto"]
    
    if 'secoes' in st.session_state:
        secoes = st.session_state['secoes']
        tipo_doc = st.session_state['tipo_documento']
        norma = st.session_state["config"]["norma"]
  
        # Exibir configurações atuais
        st.sidebar.markdown(f"**Configuração Atual:**")
        st.sidebar.markdown(f"- Norma: {norma}")
        st.sidebar.markdown(f"- Tipo: {tipo_doc}")
        
        # Editor de seções
        # for titulo, texto in secoes.items():
        #     with st.expander(f"📄 {titulo}", expanded=True):
        #         edited_text = st.text_area(
        #             titulo,
        #             value=texto,
        #             height=300,
        #             key=f"edit_{titulo}"
        #         )
        #         secoes[titulo] = edited_text
                
        #         # Análise completa
        #         resultados = avaliar_secao_completa(edited_text, titulo)
        #         exibir_resultados_avaliacao(resultados, titulo)
                
        #         # Estatísticas
        #         palavras = contar_palavras_relevantes(edited_text)
        #         citacoes, _ = verificar_citacoes_completas(edited_text)
        #         st.caption(f"📊 Estatísticas: {palavras} palavras relevantes | {citacoes} citações válidas")
        for titulo, texto in secoes.items():
            with st.expander(f"📄 {titulo}", expanded=True):
                edited_text = st.text_area(
                    titulo,
                    value=texto,
                    height=300,
                    key=f"edit_{titulo}"
                )
                secoes[titulo] = edited_text
                
                # Análise completa
                resultados = avaliar_secao_completa(edited_text, titulo)
                exibir_resultados_avaliacao(resultados, titulo, edited_text)  # Agora passamos o texto atual
        # Controles gerais
        st.divider()
        col1, col2, col3 = st.columns(3)
        
        with col1:
            novo_nome = st.text_input("Nome do Projeto", value=st.session_state['nome_projeto'])
            if novo_nome != st.session_state['nome_projeto']:
                st.session_state['nome_projeto'] = novo_nome
        
        with col2:
            if st.button("💾 Salvar Projeto", use_container_width=True):
                relatorio = str(resultados)  # Simplificado para exemplo
                if save_to_db(st.session_state['nome_projeto'], secoes, relatorio, tipo_doc, norma):
                    st.success("Projeto salvo com sucesso!")
        
        with col3:
            doc_bytes = gerar_documento_editado(secoes)
            st.download_button(
                label="⬇️ Baixar DOCX",
                data=doc_bytes.getvalue(),
                file_name=f"{st.session_state['nome_projeto']}_editado.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        
        # Análise global com Gemini
        st.divider()
        if st.button("🧠 Gerar Análise Completa com IA"):
            with st.spinner("Analisando documento..."):
                try:
                    prompt = f"""Analise este documento acadêmico do tipo {tipo_doc} (norma {norma}):
                    - Coerência global
                    - Qualidade argumentativa
                    - Adequação às normas {norma}
                    - Sugestões de melhoria
                    
                    Documento:
                    {secoes}"""
                    
                    response = model.generate_content(prompt)
                    st.success("Análise da IA:")
                    st.write(response.text)
                except Exception as e:
                    st.error(f"Erro na análise: {e}")

if __name__ == "__main__":
    main()